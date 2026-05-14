from __future__ import annotations

import base64
import csv
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import Literal

import pandas as pd
from docx import Document
from fastapi import HTTPException, UploadFile, status
from openpyxl import load_workbook

from app.config import settings

DocumentKind = Literal["image", "pdf", "text"]

IMAGE_MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}

TEXT_EXTENSIONS = {".csv", ".txt", ".md", ".json"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xls"}
DOCUMENT_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
ALLOWED_EXTENSIONS = (
    set(IMAGE_MEDIA_TYPES) | TEXT_EXTENSIONS | SPREADSHEET_EXTENSIONS | DOCUMENT_EXTENSIONS | PDF_EXTENSIONS
)

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/webp",
    "text/plain",
    "text/csv",
    "text/markdown",
    "application/json",
    "application/octet-stream",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


@dataclass(frozen=True)
class ProcessedDocument:
    filename: str
    kind: DocumentKind
    media_type: str
    data: str | None = None
    text: str | None = None


def _normalize_mime(content_type: str | None) -> str:
    if not content_type:
        return "application/octet-stream"
    return content_type.split(";")[0].strip().lower()


def _truncate_text(text: str) -> str:
    return text[:80_000]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def _worksheet_to_csv_text(rows: list[list[object]]) -> str:
    output = StringIO()
    writer = csv.writer(output)
    writer.writerows(rows)
    return output.getvalue()


def _extract_xlsx_text(data: bytes) -> str:
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    chunks: list[str] = []

    for worksheet in workbook.worksheets:
        rows = [
            ["" if cell is None else cell for cell in row]
            for row in worksheet.iter_rows(values_only=True)
        ]
        if rows:
            chunks.append(f"Sheet: {worksheet.title}\n{_worksheet_to_csv_text(rows)}")

    return "\n\n".join(chunks)


def _extract_xls_text(data: bytes) -> str:
    sheets = pd.read_excel(BytesIO(data), sheet_name=None, dtype=str)
    chunks: list[str] = []

    for sheet_name, frame in sheets.items():
        chunks.append(f"Sheet: {sheet_name}\n{frame.fillna('').to_csv(index=False)}")

    return "\n\n".join(chunks)


async def process_upload(file: UploadFile) -> ProcessedDocument:
    filename = file.filename or "document"
    extension = Path(filename).suffix.lower()
    content_type = _normalize_mime(file.content_type)

    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file extension: {extension or 'none'}",
        )

    if content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported MIME type: {content_type}",
        )

    data = await file.read()

    if len(data) > settings.max_file_size_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File must be {settings.max_file_size_mb} MB or smaller",
        )

    if extension in IMAGE_MEDIA_TYPES:
        return ProcessedDocument(
            filename=filename,
            kind="image",
            media_type=IMAGE_MEDIA_TYPES[extension],
            data=base64.b64encode(data).decode("ascii"),
        )

    if extension in PDF_EXTENSIONS:
        return ProcessedDocument(
            filename=filename,
            kind="pdf",
            media_type="application/pdf",
            data=base64.b64encode(data).decode("ascii"),
        )

    if extension in DOCUMENT_EXTENSIONS:
        text = _extract_docx_text(data)
    elif extension == ".xlsx":
        text = _extract_xlsx_text(data)
    elif extension == ".xls":
        text = _extract_xls_text(data)
    else:
        text = _decode_text(data)

    return ProcessedDocument(
        filename=filename,
        kind="text",
        media_type=content_type,
        text=_truncate_text(text),
    )
